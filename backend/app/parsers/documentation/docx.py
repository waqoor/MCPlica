import re
from io import BytesIO
from uuid import UUID

from docx import Document
from docx.table import Table, _Cell  # pyright: ignore[reportPrivateUsage]
from docx.text.paragraph import Paragraph

from app.core.exceptions import SourceParseError

from .common import ensure_text_limit
from .models import DocumentSection, NormalizedDocument
from .office import detect_office_document_format

_HEADING = re.compile(r"^heading\s+([1-6])$", re.IGNORECASE)
_MAX_DOCX_BLOCKS = 200_000
_MAX_DOCX_TABLE_CELLS = 2_000_000


def parse_docx(
    value: bytes,
    *,
    source_version_id: UUID,
    title: str | None,
    max_text_chars: int,
) -> NormalizedDocument:
    if detect_office_document_format(value) != "docx":
        raise SourceParseError("Office document is not a DOCX document")
    try:
        document = Document(BytesIO(value))
    except Exception as exc:
        raise SourceParseError("DOCX documentation is invalid or unreadable") from exc

    metadata_title = (document.core_properties.title or "").strip() or None
    document_title = title or metadata_title
    root = document_title or "Document"
    headings: list[str] = []
    current_lines: list[str] = []
    sections: list[DocumentSection] = []
    total_chars = 0
    block_count = 0
    table_cells = 0
    table_number = 0

    def add_text(text: str) -> None:
        nonlocal total_chars
        total_chars += len(text) + 1
        if total_chars > max_text_chars:
            raise SourceParseError("DOCX extracted text exceeds the configured limit")

    def flush() -> None:
        text = "\n\n".join(current_lines).strip()
        current_lines.clear()
        if not text:
            return
        sections.append(
            DocumentSection(
                path=list(headings) or [root],
                heading=headings[-1] if headings else document_title,
                text=text,
                ordinal=len(sections),
            )
        )

    def count_block() -> None:
        nonlocal block_count
        block_count += 1
        if block_count > _MAX_DOCX_BLOCKS:
            raise SourceParseError("DOCX documentation exceeds the block limit")

    def paragraph_text(paragraph: Paragraph) -> str:
        count_block()
        return paragraph.text.strip()

    def render_cell(cell: _Cell) -> str:
        parts: list[str] = []
        for inner in cell.iter_inner_content():
            if isinstance(inner, Paragraph):
                text = paragraph_text(inner)
                if text:
                    parts.append(text)
            else:
                nested = render_table(inner)
                if nested:
                    parts.append("\\n".join(nested))
        return "\\n".join(parts)

    def render_table(table: Table) -> list[str]:
        nonlocal table_cells
        count_block()
        lines: list[str] = []
        for row in table.rows:
            table_cells += len(row.cells)
            if table_cells > _MAX_DOCX_TABLE_CELLS:
                raise SourceParseError("DOCX documentation exceeds the table cell limit")
            rendered = [render_cell(cell).replace("\n", "\\n") for cell in row.cells]
            while rendered and not rendered[-1]:
                rendered.pop()
            if rendered:
                line = "\t".join(rendered)
                add_text(line)
                lines.append(line)
        return lines

    for block in document.iter_inner_content():
        if isinstance(block, Table):
            flush()
            table_number += 1
            lines = render_table(block)
            if lines:
                heading = f"Table {table_number}"
                sections.append(
                    DocumentSection(
                        path=[*(headings or [root]), heading],
                        heading=heading,
                        text="\n".join(lines),
                        ordinal=len(sections),
                    )
                )
            continue
        paragraph = block
        text = paragraph_text(paragraph)
        if not text:
            continue
        style = paragraph.style
        style_name = (style.name or "") if style is not None else ""
        heading_match = _HEADING.fullmatch(style_name.strip())
        if heading_match:
            flush()
            level = int(heading_match.group(1))
            headings[:] = headings[: level - 1]
            headings.append(text)
            add_text(text)
        else:
            add_text(text)
            current_lines.append(text)
    flush()
    if not sections:
        raise SourceParseError("DOCX documentation contains no extractable text")
    if document_title is None and sections[0].heading:
        document_title = sections[0].path[0]
    text = "\n\n".join(
        "\n".join([*(section.path[-1:] if section.heading else []), section.text])
        for section in sections
    )
    ensure_text_limit(text, label="DOCX", max_text_chars=max_text_chars)
    return NormalizedDocument(
        source_version_id=source_version_id,
        title=document_title,
        text=text,
        sections=sections,
        metadata={
            "format": "docx",
            "paragraph_count": len(document.paragraphs),
            "table_count": len(document.tables),
        },
    )
