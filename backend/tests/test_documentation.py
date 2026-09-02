import hashlib
from io import BytesIO
from uuid import UUID

import pymupdf
import pytest
from docx import Document
from openpyxl import Workbook

from app.core.exceptions import SourceParseError
from app.parsers.documentation import chunk_document, parse_document


def test_markdown_sections_and_chunk_ids_are_deterministic() -> None:
    source_id = UUID(int=20)
    document = parse_document(
        b"# Products\n\nOverview.\n\n## Lookup\n\nUse the product identifier.",
        detected_format="markdown",
        source_version_id=source_id,
        title=None,
    )
    assert document.title == "Products"
    assert document.sections[1].path == ["Products", "Lookup"]

    project_id = UUID(int=21)
    generation_id = UUID(int=22)
    source_hash = hashlib.sha256(b"source").hexdigest()
    first = chunk_document(
        document,
        project_id=project_id,
        generation_id=generation_id,
        source_content_sha256=source_hash,
        max_chars=500,
        overlap_chars=50,
    )
    second = chunk_document(
        document,
        project_id=project_id,
        generation_id=generation_id,
        source_content_sha256=source_hash,
        max_chars=500,
        overlap_chars=50,
    )
    assert first == second
    assert first[0].chunk_id.startswith("chunk_")


def test_html_omits_active_content() -> None:
    document = parse_document(
        b"<html><head><title>Docs</title><script>secret()</script></head>"
        b"<body><h1>Safe</h1><p>Visible text</p></body></html>",
        detected_format="html",
        source_version_id=UUID(int=23),
        title=None,
    )
    assert "Visible text" in document.text
    assert "secret" not in document.text


@pytest.mark.parametrize(
    ("detected_format", "value", "expected"),
    [
        ("json", b'{"endpoint": "/widgets", "method": "GET"}', '"/widgets"'),
        ("markdown", b"# Widget API\n\nList widgets.", "List widgets"),
        ("text", b"Widget API reference", "Widget API reference"),
        (
            "csv",
            b"method,path,summary\nGET,/widgets,List widgets\n",
            "GET\t/widgets\tList widgets",
        ),
    ],
)
def test_structured_and_text_document_formats_are_normalized(
    detected_format: str,
    value: bytes,
    expected: str,
) -> None:
    document = parse_document(
        value,
        detected_format=detected_format,
        source_version_id=UUID(int=24),
        title="Widget guide",
    )
    assert expected in document.text
    assert document.metadata["format"] == detected_format


def test_xlsx_formulas_are_extracted_as_text_without_execution() -> None:
    output = BytesIO()
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Operations"
    worksheet.append(["method", "path", "description"])
    worksheet.append(["GET", "/widgets", '=CONCAT("List"," widgets")'])
    workbook.save(output)
    workbook.close()

    document = parse_document(
        output.getvalue(),
        detected_format="xlsx",
        source_version_id=UUID(int=25),
        title="Workbook guide",
    )

    assert "GET\t/widgets" in document.text
    assert '=CONCAT("List"," widgets")' in document.text
    assert document.metadata["format"] == "xlsx"


def test_docx_paragraphs_headings_and_tables_are_extracted() -> None:
    output = BytesIO()
    source = Document()
    source.add_heading("Widget API", level=1)
    source.add_paragraph("Use the listWidgets operation to retrieve widgets.")
    table = source.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Method"
    table.cell(0, 1).text = "Path"
    table.cell(1, 0).text = "GET"
    table.cell(1, 1).text = "/widgets"
    source.save(output)

    document = parse_document(
        output.getvalue(),
        detected_format="docx",
        source_version_id=UUID(int=26),
        title=None,
    )

    assert document.title == "Widget API"
    assert "listWidgets" in document.text
    assert "GET\t/widgets" in document.text
    assert document.metadata["format"] == "docx"


def test_docx_preserves_interleaved_table_order_and_heading_context() -> None:
    output = BytesIO()
    source = Document()
    source.add_heading("Widget API", level=1)
    source.add_paragraph("Before the first table.")
    first = source.add_table(rows=1, cols=2)
    first.cell(0, 0).text = "GET"
    first.cell(0, 1).text = "/widgets"
    source.add_paragraph("After the first table.")
    source.add_heading("Administration", level=2)
    second = source.add_table(rows=1, cols=2)
    second.cell(0, 0).text = "DELETE"
    second.cell(0, 1).text = "/widgets/{id}"
    source.save(output)

    document = parse_document(
        output.getvalue(),
        detected_format="docx",
        source_version_id=UUID(int=29),
        title=None,
    )

    assert [section.text for section in document.sections] == [
        "Before the first table.",
        "GET\t/widgets",
        "After the first table.",
        "DELETE\t/widgets/{id}",
    ]
    assert document.sections[1].path == ["Widget API", "Table 1"]
    assert document.sections[3].path == ["Widget API", "Administration", "Table 2"]


def test_pdf_extracts_text_with_page_provenance() -> None:
    source = pymupdf.open()
    page = source.new_page()
    page.insert_text((72, 72), "Widget PDF API reference")
    value = source.tobytes()
    source.close()

    document = parse_document(
        value,
        detected_format="pdf",
        source_version_id=UUID(int=27),
        title="PDF guide",
    )

    assert "Widget PDF API reference" in document.text
    assert document.sections[0].path == ["PDF guide", "Page 1"]


def test_document_text_limit_applies_to_non_pdf_formats() -> None:
    with pytest.raises(SourceParseError, match="configured limit"):
        parse_document(
            b"123456",
            detected_format="text",
            source_version_id=UUID(int=28),
            title=None,
            max_text_chars=5,
        )
